"""Parse resume files (PDF / DOCX) into plain text."""

from pathlib import Path

from src.utils.helpers import get_logger

logger = get_logger(__name__)


def parse_pdf(file_path: str | Path) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 is required: pip install pypdf2")

    text_parts = []
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    text = "\n".join(text_parts).strip()
    logger.info(f"[PDF] {Path(file_path).name} → {len(text)} chars")
    return text


def parse_docx(file_path: str | Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs).strip()
    logger.info(f"[DOCX] {Path(file_path).name} → {len(text)} chars")
    return text


def parse_csv(file_path: str | Path) -> list[dict]:
    """Parse a CSV file where each row = one candidate.

    Looks for a text column in order: 'text', 'resume_text', 'resume', 'content'.
    Falls back to concatenating all string columns if none found.

    Returns:
        List of {"name": str, "text": str}
    """
    import pandas as pd

    df = pd.read_csv(file_path, dtype=str).fillna("")

    # Find the text column
    text_col = None
    for candidate_col in ["text", "resume_text", "resume", "content"]:
        if candidate_col in df.columns:
            text_col = candidate_col
            break

    # Find the name column
    name_col = None
    for candidate_col in ["name", "candidate_name", "candidate", "filename"]:
        if candidate_col in df.columns:
            name_col = candidate_col
            break

    records = []
    for i, row in df.iterrows():
        if text_col:
            text = row[text_col].strip()
        else:
            # Concatenate all columns as fallback
            text = " ".join(str(v) for v in row.values if str(v).strip()).strip()

        name = row[name_col].strip() if name_col else f"row_{i+1}"
        if text:
            records.append({"name": name, "text": text})

    logger.info(f"[CSV] {Path(file_path).name} → {len(records)} candidates")
    return records


def parse_resume(file_path: str | Path) -> str:
    """Auto-detect file type and return plain text.

    Supports: .pdf, .docx, .txt
    For .csv use parse_csv() instead (returns multiple candidates).
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return parse_pdf(file_path)
    elif suffix == ".docx":
        return parse_docx(file_path)
    elif suffix == ".txt":
        text = file_path.read_text(encoding="utf-8").strip()
        logger.info(f"[TXT] {file_path.name} → {len(text)} chars")
        return text
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Use .pdf, .docx, or .txt")


def parse_uploaded_file(uploaded_file) -> list[dict]:
    """Parse a Streamlit UploadedFile object.

    Returns list of {"name": str, "text": str} regardless of file type.
    CSV may return multiple candidates; PDF/DOCX/TXT returns one.
    """
    import io
    import pandas as pd

    name = uploaded_file.name
    suffix = Path(name).suffix.lower()
    data = uploaded_file.read()

    if suffix == ".pdf":
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        text = "\n".join(
            p.extract_text() for p in reader.pages if p.extract_text()
        ).strip()
        return [{"name": name, "text": text}]

    elif suffix == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [{"name": name, "text": text}]

    elif suffix == ".txt":
        text = data.decode("utf-8", errors="ignore").strip()
        return [{"name": name, "text": text}]

    elif suffix == ".csv":
        df = pd.read_csv(io.BytesIO(data), dtype=str).fillna("")
        text_col = next(
            (c for c in ["text", "resume_text", "resume", "content"] if c in df.columns),
            None,
        )
        name_col = next(
            (c for c in ["name", "candidate_name", "candidate"] if c in df.columns),
            None,
        )
        records = []
        for i, row in df.iterrows():
            text = row[text_col].strip() if text_col else " ".join(
                str(v) for v in row.values if str(v).strip()
            )
            cname = row[name_col].strip() if name_col else f"{name}::row{i+1}"
            if text:
                records.append({"name": cname, "text": text})
        return records

    else:
        raise ValueError(f"Unsupported file type: {suffix}")
